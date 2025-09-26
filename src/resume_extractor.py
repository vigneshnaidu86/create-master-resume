#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Resume Extractor: scans INPUT_DIR for PDF/DOCX, extracts text, calls LLM to map to JSON, and writes per-file JSONs.
Optionally merges all JSONs into a master resume via merge_jsons.py.
"""

from __future__ import annotations
import os, json, re, requests, sys
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime

from config import INPUT_DIR, OUTPUT_DIR, PROVIDER, MODEL, OLLAMA_HOST, OPENAI_API_KEY, MERGE_AFTER_EXTRACTION

# Text extraction backends
import fitz  # PyMuPDF (default)
from docx import Document

# -----------------------------
# JSON schema example
# -----------------------------
JSON_SCHEMA_EXAMPLE = {
  "basics": {"name": "", "label": "", "image": "", "email": "", "phone": "", "url": "", "summary": "",
    "location": {"address": "", "postalCode": "", "city": "", "countryCode": "", "region": ""},
    "profiles": [{"network": "", "username": "", "url": ""}]},
  "work": [{"name": "", "position": "", "url": "", "startDate": "", "endDate": "", "summary": "", "highlights": []}],
  "volunteer": [{"organization": "", "position": "", "url": "", "startDate": "", "endDate": "", "summary": "", "highlights": []}],
  "education": [{"institution": "", "url": "", "area": "", "studyType": "", "startDate": "", "endDate": "", "score": "", "courses": []}],
  "awards": [{"title": "", "date": "", "awarder": "", "summary": ""}],
  "certificates": [{"name": "", "date": "", "issuer": "", "url": ""}],
  "publications": [{"name": "", "publisher": "", "releaseDate": "", "url": "", "summary": ""}],
  "skills": [{"name": "", "level": "", "keywords": []}],
  "languages": [{"language": "", "fluency": ""}],
  "interests": [{"name": "", "keywords": []}],
  "references": [{"name": "", "reference": ""}],
  "projects": [{"name": "", "startDate": "", "endDate": "", "description": "", "highlights": [], "url": ""}]
}

EXTRACTION_SYSTEM_PROMPT = """You are a careful information extraction assistant.
You will be given the full text of a resume. Your job is to convert it into a STRICT JSON object following the provided JSON schema exactly (keys, nesting, and arrays).
- Use ISO dates (YYYY-MM-DD) where possible; if only year/month known, use YYYY-MM-01.
- If a field is unknown, return an empty string or an empty array/object (do not invent).
- For skills: group related keywords under a reasonable 'name' and include keywords list; set 'level' only if explicitly stated.
- For work: include concise bullet-style 'highlights' from achievements/impact/scale/metrics.
- Do NOT include any commentary—ONLY return valid JSON.
"""

EXTRACTION_USER_PROMPT_TEMPLATE = """JSON schema (example; keep the exact structure/keys, values should reflect the resume):
{schema}

Resume text:
\"\"\"
{resume_text}
\"\"\"
"""

# -----------------------------
# File discovery & reading
# -----------------------------

def find_resume_files(root: Path) -> List[Path]:
    all_files = list(root.rglob("*.pdf")) + list(root.rglob("*.docx"))
    buckets = {}
    for f in all_files:
        key = (f.parent, f.stem)
        buckets.setdefault(key, {})
        buckets[key][f.suffix.lower()] = f
    chosen = []
    for _, variants in buckets.items():
        if ".docx" in variants:
            chosen.append(variants[".docx"])
        else:
            chosen.append(variants.get(".pdf"))
    return [c for c in chosen if c is not None]

def read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join([p for p in parts if p and p.strip()])

def read_pdf_text(path: Path) -> str:
    text_chunks = []
    with fitz.open(path) as doc:
        for page in doc:
            text_chunks.append(page.get_text(\"text\"))
    raw = \"\n\".join(text_chunks)
    raw = re.sub(r\"[ \\t]+\", \" \", raw)
    raw = re.sub(r\"\\n{3,}\", \"\\n\\n\", raw)
    return raw.strip()

def extract_text(path: Path) -> str:
    if path.suffix.lower() == \".docx\":
        return read_docx_text(path)
    elif path.suffix.lower() == \".pdf\":
        return read_pdf_text(path)
    else:
        raise ValueError(f\"Unsupported file type: {path.suffix}\")

# -----------------------------
# LLM callers
# -----------------------------

def safe_json_loads(s: str) -> Dict[str, Any]:
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        last_brace = s.rfind(\"}\")
        if last_brace != -1:
            s2 = s[:last_brace+1]
            return json.loads(s2)
        raise

def call_ollama(resume_text: str, model: str = \"llama3.2\", host: Optional[str] = None) -> Dict[str, Any]:
    base = host or os.environ.get(\"OLLAMA_HOST\", \"http://localhost:11434\")
    url = f\"{base}/api/chat\"
    payload = {
        \"model\": model,
        \"messages\": [
            {\"role\": \"system\", \"content\": EXTRACTION_SYSTEM_PROMPT},
            {\"role\": \"user\", \"content\": EXTRACTION_USER_PROMPT_TEMPLATE.format(
                schema=json.dumps(JSON_SCHEMA_EXAMPLE, indent=2),
                resume_text=resume_text[:120000]
            )}
        ],
        \"format\": \"json\",
        \"stream\": False
    }
    r = requests.post(url, json=payload, timeout=300)
    r.raise_for_status()
    data = r.json()
    content = data.get(\"message\", {}).get(\"content\", \"\").strip()
    return safe_json_loads(content)

def call_openai(resume_text: str, model: str = \"gpt-4o-mini\") -> Dict[str, Any]:
    api_key = OPENAI_API_KEY or os.environ.get(\"OPENAI_API_KEY\")
    if not api_key:
        raise RuntimeError(\"OPENAI_API_KEY not set (env or config)\" )
    url = \"https://api.openai.com/v1/chat/completions\"
    headers = {\"Authorization\": f\"Bearer {api_key}\", \"Content-Type\": \"application/json\"}
    payload = {
        \"model\": model,
        \"temperature\": 0,
        \"response_format\": {\"type\": \"json_object\"},
        \"messages\": [
            {\"role\": \"system\", \"content\": EXTRACTION_SYSTEM_PROMPT},
            {\"role\": \"user\", \"content\": EXTRACTION_USER_PROMPT_TEMPLATE.format(
                schema=json.dumps(JSON_SCHEMA_EXAMPLE, indent=2),
                resume_text=resume_text[:120000]
            )}
        ]
    }
    r = requests.post(url, headers=headers, json=payload, timeout=300)
    r.raise_for_status()
    data = r.json()
    content = data[\"choices\"][0][\"message\"][\"content\"].strip()
    return safe_json_loads(content)

def extract_one_resume_to_json(path: Path, provider: str, model: str, ollama_host: Optional[str] = None) -> Dict[str, Any]:
    text = extract_text(path)
    if not text.strip():
        raise RuntimeError(f\"No text extracted from {path}\")
    if provider == \"ollama\":
        return call_ollama(text, model=model, host=ollama_host)
    elif provider == \"openai\":
        return call_openai(text, model=model)
    else:
        raise ValueError(\"provider must be 'ollama' or 'openai'\")

# -----------------------------
# Pipeline
# -----------------------------

def run_pipeline():
    input_dir = Path(INPUT_DIR).expanduser().resolve()
    out_dir = Path(OUTPUT_DIR).expanduser().resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    files = find_resume_files(input_dir)
    if not files:
        print(\"No resumes found.\")
        sys.exit(1)

    extracted_jsons: List[Dict[str, Any]] = []
    for f in files:
        print(f\"[+] Processing: {f}\")
        try:
            data = extract_one_resume_to_json(f, provider=PROVIDER, model=MODEL, ollama_host=OLLAMA_HOST)
            per_file_out = out_dir / (f.stem + \".json\")
            with per_file_out.open(\"w\", encoding=\"utf-8\") as fh:
                json.dump(data, fh, ensure_ascii=False, indent=2)
            extracted_jsons.append(data)
        except Exception as e:
            print(f\"[!] Failed on {f}: {e}\")

    # Optional merge step
    if MERGE_AFTER_EXTRACTION:
        try:
            from merge_jsons import merge_all_resumes
            master = merge_all_resumes(extracted_jsons)
            master_path = out_dir / \"master_resume.json\"
            with master_path.open(\"w\", encoding=\"utf-8\") as fh:
                json.dump(master, fh, ensure_ascii=False, indent=2)
            print(f\"[✓] Master resume written to: {master_path}\")
        except Exception as e:
            print(f\"[!] Merge step failed. You can run 'python src/merge_jsons.py' later. Error: {e}\")

if __name__ == \"__main__\":
    run_pipeline()
